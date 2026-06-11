"""
word_export.py — python-docx 기반 Word 문서 생성 (표 형식)
"""

import io
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "맑은 고딕"

SECTION_LABELS = [
    ("section_01",  "1. 사업 개요"),
    ("section_02",  "2. 위험성 분석"),
    ("section_law", "3. 관련 법령 검토"),
    ("section_04",  "4. 타 지자체 유사 사례"),
    ("section_08",  "5. 의회 예상 질의·대응논리"),
    ("section_09",  "6. 종합 검토의견"),
]

PRECHECK_ITEMS = [
    ("보안성 검토",      "개인정보 처리 포함 시"),
    ("개인정보 영향평가", "5만명 이상 처리 시"),
    ("지방재정투자심사",  "총사업비 10억 이상 시"),
    ("의회 의결",        "공유재산 취득·처분 시"),
    ("법령 저촉 여부",   "모든 사업 확인 필요"),
    ("타 부서 협의",     "관련 부서 협의 필요"),
    ("환경영향평가",     "개발·시설 사업 해당 시"),
    ("안전영향평가",     "다중이용시설 포함 시"),
]

RISK_LEVEL_MAP = {
    "고": "고위험 (재검토/부적정)",
    "중": "중위험 (조건부 적정)",
    "저": "저위험 (적정)",
}


def _set_font(run, size: int = 10, bold: bool = False, color: RGBColor = None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _cell_shading(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, size=11, bold=True, color=RGBColor(0x00, 0x30, 0x87))


def _parse_pipe_table(text: str) -> list[list[str]]:
    """파이프(|) 구분 텍스트 → 행/열 리스트. 마크다운 구분선 제거."""
    rows = []
    for line in text.strip().splitlines():
        s = line.strip()
        if not s:
            continue
        if re.match(r'^[\s|:\-]+$', s):   # |---|---| 구분선 제거
            continue
        s = s.strip("|").strip()
        cols = [c.strip() for c in s.split("|")]
        cols = [c for c in cols if c != ""]
        if cols:
            rows.append(cols)
    return rows


def _add_hyperlink_to_cell(cell, url: str, display: str = "바로가기"):
    """Word 셀에 하이퍼링크 추가 (XML 직접 조작)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    wr = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")   # 9pt = 18 half-points
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    wr.append(rPr)

    t = OxmlElement("w:t")
    t.text = display
    wr.append(t)

    hyperlink.append(wr)
    para._p.append(hyperlink)


def _add_pipe_table(doc: Document, text: str):
    """파이프 구분 텍스트를 Word 표로 추가. URL 셀은 하이퍼링크로 처리."""
    rows = _parse_pipe_table(text)
    if not rows:
        p = doc.add_paragraph()
        p.add_run("(내용 없음)").font.name = FONT_NAME
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            is_header = r_idx == 0

            if not is_header and (value.startswith("https://") or value.startswith("http://")):
                _add_hyperlink_to_cell(cell, value, "바로가기")
            else:
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(value)
                _set_font(run, size=9, bold=is_header)

            if is_header:
                _cell_shading(cell, "DCE6F1")


def _add_precheck_table(doc: Document, section_03_text: str):
    """사전 검토 체크리스트 표를 Word로 추가."""
    headers = ["검토 항목", "적용 조건", "AI 예상 판단", "판단 근거"]
    VALID = {"해당", "확인 필요", "해당 없음"}

    # AI 출력 파싱
    ai_lines = []
    if section_03_text:
        for line in section_03_text.strip().splitlines():
            s = line.strip()
            if not s or re.match(r'^[\s|:\-]+$', s):
                continue
            s = s.strip("|").strip()
            parts = [c.strip() for c in s.split("|")]
            parts = [c for c in parts if c]
            ai_lines.append(parts)

    rows_data = [headers]
    for i, (name, cond) in enumerate(PRECHECK_ITEMS):
        if i < len(ai_lines):
            p = ai_lines[i]
            judgment = p[0] if p and p[0] in VALID else "확인 필요"
            reason = p[1] if len(p) > 1 else ""
        else:
            judgment = "확인 필요"
            reason = ""
        rows_data.append([name, cond, judgment, reason])

    table = doc.add_table(rows=len(rows_data), cols=4)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows_data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            is_header = r_idx == 0
            _set_font(run, size=9, bold=is_header)
            if is_header:
                _cell_shading(cell, "DCE6F1")
            elif c_idx == 2:  # AI 판단 열 색상
                if value == "해당":
                    _cell_shading(cell, "FFE6E6")
                elif value == "해당 없음":
                    _cell_shading(cell, "E6FFE6")


def _add_info_table(doc: Document, report: dict):
    budget = report.get("budget")
    budget_str = f"{int(budget):,}원" if budget else "미기재"
    rows = [
        ("사업명", report.get("project_name", "")),
        ("사업 유형", report.get("project_type", "")),
        ("주요 대상", report.get("target") or "미기재"),
        ("총 예산", budget_str),
        ("사업 기간", report.get("period") or "미기재"),
        ("AI 위험도", RISK_LEVEL_MAP.get(report.get("risk_level", "중"), "중위험")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        lc = table.cell(i, 0)
        vc = table.cell(i, 1)
        lc.width = Cm(3.5)
        _cell_shading(lc, "F2F2F2")
        for cell, text, bold in [(lc, label, True), (vc, value, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            _set_font(run, size=10, bold=bold)


def _add_section_body(doc: Document, text: str):
    """섹션 콘텐츠를 Word로 추가. 빈 줄로 구분된 복수 블록 지원 (캡션 + 표)."""
    if not text:
        p = doc.add_paragraph()
        p.add_run("(내용 없음)").font.name = FONT_NAME
        return

    # 빈 줄로 블록 분리
    blocks: list[list[str]] = []
    current: list[str] = []
    for raw_line in text.splitlines():
        s = raw_line.strip()
        if not s:
            if current:
                blocks.append(current)
                current = []
        else:
            if re.match(r'^[\s|:\-]+$', s):
                continue
            current.append(s)
    if current:
        blocks.append(current)

    for block in blocks:
        _add_block(doc, block)


def _add_block(doc: Document, lines: list[str]):
    """단일 블록을 Word로 추가. [캡션] + 표 또는 일반 텍스트."""
    caption = ""
    table_lines = []

    for line in lines:
        if line.startswith("[") and line.endswith("]") and "|" not in line:
            caption = line[1:-1]
        else:
            table_lines.append(line)

    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(caption)
        _set_font(run, size=9, bold=True, color=RGBColor(0x33, 0x33, 0x33))

    if table_lines and "|" in table_lines[0]:
        _add_pipe_table(doc, "\n".join(table_lines))
    elif table_lines:
        p = doc.add_paragraph()
        run = p.add_run("\n".join(table_lines))
        _set_font(run, size=9)


def generate_word_bytes(report: dict) -> bytes:
    doc = Document()

    # A4 여백
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # 제목
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    _set_font(title_p.add_run("정책사업 사전검토 보고서"), size=16, bold=True)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    _set_font(sub_p.add_run(report.get("project_name", "")), size=13, bold=True)

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.paragraph_format.space_after = Pt(16)
    _set_font(org_p.add_run("고양특례시"), size=11)

    # 기본 정보 표
    _add_info_table(doc, report)
    doc.add_paragraph()

    # 사전 검토 체크리스트
    _add_section_heading(doc, "① 사전 검토 체크리스트")
    _add_precheck_table(doc, report.get("section_03", ""))
    doc.add_paragraph()

    # 5개 보고서 섹션
    for key, label in SECTION_LABELS:
        content = report.get(key, "")
        _add_section_heading(doc, label)
        _add_section_body(doc, content)
        doc.add_paragraph()

    # 푸터
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(12)
    run = footer_p.add_run(
        "※ 본 보고서는 Claude AI가 자동 생성한 내부 검토 초안입니다. "
        "법령 조문 번호와 타 지자체 사례는 담당자가 직접 확인 후 활용하십시오."
    )
    _set_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
